"""
Create a sample DOCX template with all monitoring form fields as placeholders.
This script generates a comprehensive template that users can customize.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Create a new Document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Add header with company name
header = doc.add_heading('SGV SUPER SECURITY SERVICE PVT. LTD', 0)
header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
header_run = header.runs[0]
header_run.font.color.rgb = RGBColor(26, 58, 107)

# Add report title
title = doc.add_heading('VIGILANCE & MONITORING INSPECTION REPORT', level=1)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()  # Blank line

# BASIC INFORMATION SECTION
doc.add_heading('BASIC INFORMATION', level=2)
table1 = doc.add_table(rows=7, cols=2)
table1.style = 'Light Grid Accent 1'

basic_info = [
    ('Date:', '{Date}'),
    ('Time:', '{Time}'),
    ('Timestamp:', '{Timestamp}'),
    ('Site Name:', '{Site Name}'),
    ('Shift:', '{Shift}'),
    ('Inspected By:', '{Inspected By}'),
    ('Email Address:', '{Email Address}')
]

for i, (label, placeholder) in enumerate(basic_info):
    table1.rows[i].cells[0].text = label
    table1.rows[i].cells[1].text = placeholder

doc.add_paragraph()

# DOCUMENTATION CHECKS SECTION
doc.add_heading('DOCUMENTATION CHECKS', level=2)
table2 = doc.add_table(rows=5, cols=2)
table2.style = 'Light Grid Accent 1'

doc_checks = [
    ('Attendance Register', '{Documentation Check [Attendance Register]}'),
    ('Handling / Taking Over Register', '{Documentation Check [Handling / Taking Over Register]}'),
    ('Visitor Log Register', '{Documentation Check [Visitor Log Register]}'),
    ('Incident Log', '{Documentation Check [Incident Log]}'),
    ('Overall Documentation Status', '{Performance Check [Row 5]}')
]

for i, (label, placeholder) in enumerate(doc_checks):
    table2.rows[i].cells[0].text = label
    table2.rows[i].cells[1].text = placeholder

doc.add_paragraph()

# PERFORMANCE CHECKS SECTION
doc.add_heading('PERFORMANCE ASSESSMENT', level=2)
table3 = doc.add_table(rows=4, cols=2)
table3.style = 'Light Grid Accent 1'

performance_checks = [
    ('Grooming', '{Performance Check [Grooming]}'),
    ('Alertness', '{Performance Check [Alertness]}'),
    ('Post Discipline', '{Performance Check [Post Discipline]}'),
    ('Job Awareness', '{Performance Check [Job Awareness]}')
]

for i, (label, placeholder) in enumerate(performance_checks):
    table3.rows[i].cells[0].text = label
    table3.rows[i].cells[1].text = placeholder

doc.add_paragraph()

# INCIDENT REPORTING SECTION
doc.add_heading('INCIDENT REPORTING', level=2)
p1 = doc.add_paragraph()
p1.add_run('Incident Reported (if any):').bold = True
doc.add_paragraph('{Incident Reported, if any during the QC/ Night Check (Provide Details)}')

doc.add_paragraph()
p2 = doc.add_paragraph()
p2.add_run('Action Taken:').bold = True
doc.add_paragraph('{Action Taken}')

doc.add_paragraph()

# EMPLOYEE INFORMATION SECTION
doc.add_heading('EMPLOYEE INFORMATION', level=2)
p3 = doc.add_paragraph()
p3.add_run('List of Employees on Duty:').bold = True
doc.add_paragraph('{List of Employees (Please mention Emp Id, Name and Father Name)}')

doc.add_paragraph()

# EMPLOYEE CASES SECTION
doc.add_heading('EMPLOYEE CASES FOUND', level=2)
table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Light Grid Accent 1'

employee_cases = [
    ('Sleeping Cases', '{Sleeping cases Found (Provide Name, Emp Id / Father Name)}'),
    ('Employees Found Sleeping', '{Employees Found Sleeping}'),
    ('Not on Duty/Post Cases', '{Not on Duty/Post Cases Found (Provide Name, Emp Id / Father Name)}'),
    ('Misbehaving / Drunk Cases', '{Found Misbehaving / Drunk (Provide Name, Emp Id / Father Name)}'),
    ('Other Cases', '{Any other Cases Found (Provide Name, Emp Id / Father Name)}')
]

for i, (label, placeholder) in enumerate(employee_cases):
    table4.rows[i].cells[0].text = label
    table4.rows[i].cells[1].text = placeholder

doc.add_paragraph()

# ESCALATION STATUS
p4 = doc.add_paragraph()
p4.add_run('Cases Escalation Status:').bold = True
doc.add_paragraph('{Were the identified cases shared with Supervisor, FO and Manager Operations for necessary action}')

doc.add_paragraph()

# SECURITY OBSERVATIONS SECTION
doc.add_heading('SECURITY OBSERVATIONS', level=2)
p5 = doc.add_paragraph()
p5.add_run('Additional Observations:').bold = True
doc.add_paragraph('{Any other security related observations during Quality Check}')

doc.add_paragraph()

# IMAGES SECTION
doc.add_heading('INSPECTION IMAGES', level=2)
p6 = doc.add_paragraph()
p6.add_run('Image Reference:').bold = True
doc.add_paragraph('{Images}')
doc.add_paragraph()
p_note = doc.add_paragraph('Note: Image URLs/links will be displayed here. Automatic image embedding is not currently supported.')
p_note_run = p_note.runs[0]
p_note_run.font.size = Pt(9)
p_note_run.font.italic = True

doc.add_paragraph()
doc.add_paragraph()

# FOOTER
footer_para = doc.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
footer_run = footer_para.add_run('Generated by SGV Vigilance Report Generator')
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

# Save the document
output_path = 'templates/client_report_template.docx'
doc.save(output_path)

print(f"✅ DOCX template created successfully at: {output_path}")
print(f"📄 Template includes all {len(basic_info) + len(doc_checks) + len(performance_checks) + 2 + 1 + len(employee_cases) + 1 + 1 + 1} placeholders")
print("\nPlaceholder format: {Column Name}")
print("You can now customize this template and use it for client reports!")
