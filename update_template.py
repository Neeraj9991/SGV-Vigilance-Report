"""
Update the client report template to remove the note about image embedding not being supported
since we now support it.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Load existing template
template_path = 'templates/client_report_template.docx'
doc = Document(template_path)

# Find and remove the note paragraph about images not being supported
paragraphs_to_remove = []
for i, paragraph in enumerate(doc.paragraphs):
    if 'Note: Image URLs/links will be displayed here' in paragraph.text or \
       'Automatic image embedding is not currently supported' in paragraph.text:
        paragraphs_to_remove.append(paragraph)

# Remove the paragraphs
for para in paragraphs_to_remove:
    p = para._element
    p.getparent().remove(p)

# Add updated note
for i, paragraph in enumerate(doc.paragraphs):
    if 'Image Reference:' in paragraph.text:
        # Find next paragraph after "Image Reference:"
        if i + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[i + 1]
            # Update it with helpful text
            next_para.text = ''  # Clear it
            run = next_para.add_run('Images from inspection will be automatically embedded here when you generate the report.')
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            break

# Save updated template
doc.save(template_path)

print(f"✅ Updated template at: {template_path}")
print("📸 Image embedding is now enabled - images will be automatically downloaded and embedded!")
