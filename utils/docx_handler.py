"""
DOCX Template Handler Module
Handles DOCX template processing, placeholder replacement, image embedding, and PDF conversion.
"""

import os
import re
import tempfile
from io import BytesIO
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import pandas as pd


class DocxHandler:
    """Handler for DOCX template processing and PDF conversion."""
    
    def __init__(self, image_handler=None):
        """
        Initialize the DOCX handler.
        
        Args:
            image_handler: Optional ImageHandler instance for downloading images
        """
        self.temp_dir = tempfile.gettempdir()
        self.image_handler = image_handler
    
    def extract_placeholders(self, docx_file):
        """
        Extract all placeholders from a DOCX template.
        
        Args:
            docx_file: File-like object or path to DOCX file
            
        Returns:
            set: Set of placeholder names (without curly braces)
        """
        try:
            doc = Document(docx_file)
            placeholders = set()
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                matches = re.findall(r'\{([^}]+)\}', paragraph.text)
                placeholders.update(matches)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        matches = re.findall(r'\{([^}]+)\}', cell.text)
                        placeholders.update(matches)
            
            return placeholders
        
        except Exception as e:
            raise Exception(f"Error extracting placeholders: {str(e)}")
    
    def match_placeholders_to_columns(self, placeholders, df_columns):
        """
        Match placeholders to DataFrame columns.
        
        Args:
            placeholders (set): Set of placeholder names
            df_columns (list): List of DataFrame column names
            
        Returns:
            tuple: (matched_dict, unmatched_list)
                - matched_dict: {placeholder: column_name}
                - unmatched_list: [placeholder, ...]
        """
        matched = {}
        unmatched = []
        
        for placeholder in placeholders:
            # Try exact match first
            if placeholder in df_columns:
                matched[placeholder] = placeholder
            # Try case-insensitive match
            else:
                found = False
                for col in df_columns:
                    if placeholder.lower() == col.lower():
                        matched[placeholder] = col
                        found = True
                        break
                
                if not found:
                    unmatched.append(placeholder)
        
        return matched, unmatched
    
    def replace_placeholders(self, doc, data_dict, skip_images=True):
        """
        Replace placeholders in document with actual data.
        
        Args:
            doc: Document object
            data_dict: Dictionary mapping placeholder names to values
            skip_images: If True, skip 'Images' placeholder (to be handled by embed_images)
            
        Returns:
            Document: Modified document
        """
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in data_dict.items():
                # Skip image placeholders (handled by embed_images)
                if skip_images and placeholder in ['Images', 'EVIDENCE & ATTACHMENTS - Photos']:
                    continue
                    
                placeholder_text = f'{{{placeholder}}}'
                if placeholder_text in paragraph.text:
                    # Convert value to string, handle None/NaN
                    str_value = str(value) if pd.notna(value) else ''
                    paragraph.text = paragraph.text.replace(placeholder_text, str_value)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in data_dict.items():
                        # Skip image placeholders (handled by embed_images)
                        if skip_images and placeholder in ['Images', 'EVIDENCE & ATTACHMENTS - Photos']:
                            continue
                            
                        placeholder_text = f'{{{placeholder}}}'
                        if placeholder_text in cell.text:
                            str_value = str(value) if pd.notna(value) else ''
                            cell.text = cell.text.replace(placeholder_text, str_value)
        
        return doc
    
    def embed_images(self, doc, data_dict):
        """
        Find and embed images in the document, replacing image placeholders.
        
        Args:
            doc: Document object
            data_dict: Dictionary mapping placeholder names to values
            
        Returns:
            Document: Modified document with embedded images
        """
        if not self.image_handler:
            print("DEBUG: No image_handler available")
            return doc
        
        # Look for image placeholders - support both monitoring and incident reports
        # For monitoring: {Images}
        # For incidents: {EVIDENCE & ATTACHMENTS - Photos}
        images_placeholder = None
        images_value = None
        
        # Check for incident photo placeholder first
        if '{EVIDENCE & ATTACHMENTS - Photos}' in str(doc):
            images_placeholder = '{EVIDENCE & ATTACHMENTS - Photos}'
            images_value = data_dict.get('EVIDENCE & ATTACHMENTS - Photos', '')
        # Check for monitoring Images placeholder
        elif '{Images}' in str(doc):
            images_placeholder = '{Images}'
            images_value = data_dict.get('Images', '')
        
        if not images_placeholder:
            print("DEBUG: No image placeholder found")
            return doc
        
        print(f"DEBUG: embed_images called, placeholder: {images_placeholder}, images_value: {images_value}")
        
        if not images_value or pd.isna(images_value):
            print("DEBUG: No images_value or is NaN")
            return doc
        
        # Download and encode images
        print(f"DEBUG: Downloading images from: {str(images_value)}")
        image_data_list = self.image_handler.download_and_encode_images(str(images_value))
        
        print(f"DEBUG: Downloaded {len(image_data_list) if image_data_list else 0} images")
        
        if not image_data_list:
            print("DEBUG: No images downloaded")
            return doc
        
        # Find and replace images placeholder in paragraphs
        found_placeholder = False
        for paragraph in doc.paragraphs:
            if images_placeholder in paragraph.text:
                print(f"DEBUG: Found {images_placeholder} in paragraph")
                found_placeholder = True
                # Clear the placeholder text
                paragraph.text = paragraph.text.replace(images_placeholder, '')
                
                # Add images after this paragraph
                parent = paragraph._element.getparent()
                index = list(parent).index(paragraph._element) + 1
                
                for i, img_data in enumerate(image_data_list):
                    # Create a new paragraph for each image
                    new_para = doc.add_paragraph()
                    
                    # Save image to temp file
                    import base64
                    import uuid
                    temp_image_path = os.path.join(self.temp_dir, f'temp_image_{uuid.uuid4().hex}.jpg')
                    
                    try:
                        # Decode and save image
                        img_bytes = base64.b64decode(img_data['base64'])
                        with open(temp_image_path, 'wb') as f:
                            f.write(img_bytes)
                        
                        # Add image to the new paragraph
                        run = new_para.add_run()
                        run.add_picture(temp_image_path, width=Inches(5.0))
                        
                        # Add caption
                        caption_para = doc.add_paragraph(f'Image {i+1}')
                        caption_para.style = 'Caption'
                        
                        print(f"DEBUG: Embedded image {i+1} in paragraph")
                        
                    except Exception as e:
                        print(f"Warning: Could not embed image {i+1}: {str(e)}")
                    
                    finally:
                        # Cleanup temp image file
                        if os.path.exists(temp_image_path):
                            os.remove(temp_image_path)
                
                # Only process first occurrence
                break
        
        if not found_placeholder:
            print(f"DEBUG: {images_placeholder} not found in paragraphs, checking tables...")
        
        # Also check in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if images_placeholder in cell.text:
                        print(f"DEBUG: Found {images_placeholder} in table cell")
                        found_placeholder = True
                        # Clear placeholder
                        cell.text = cell.text.replace(images_placeholder, '')
                        
                        # Add images to cell
                        for i, img_data in enumerate(image_data_list):
                            import base64
                            import uuid
                            temp_image_path = os.path.join(self.temp_dir, f'temp_image_{uuid.uuid4().hex}.jpg')
                            
                            try:
                                img_bytes = base64.b64decode(img_data['base64'])
                                with open(temp_image_path, 'wb') as f:
                                    f.write(img_bytes)
                                
                                # Add paragraph with image in cell
                                para = cell.add_paragraph()
                                run = para.add_run()
                                run.add_picture(temp_image_path, width=Inches(3.0))
                                
                                print(f"DEBUG: Embedded image {i+1} in table cell")
                                
                            except Exception as e:
                                print(f"Warning: Could not embed image {i+1} in table: {str(e)}")
                            
                            finally:
                                if os.path.exists(temp_image_path):
                                    os.remove(temp_image_path)
                        
                        # Only process first occurrence
                        return doc
        
        if not found_placeholder:
            print(f"WARNING: {images_placeholder} placeholder not found anywhere in document!")
        
        return doc
    
    def generate_from_template(self, template_file, df, row_index=0):
        """
        Generate a filled DOCX document from template and data.
        
        Args:
            template_file: File-like object or path to template DOCX
            df (pd.DataFrame): DataFrame containing data
            row_index (int): Index of row to use for data (default: 0, first row)
            
        Returns:
            BytesIO: Filled DOCX file as bytes
        """
        try:
            # Load template
            doc = Document(template_file)
            
            # Get data from specified row
            if len(df) == 0:
                raise ValueError("DataFrame is empty")
            
            if row_index >= len(df):
                row_index = 0
            
            # Create data dictionary from row
            data_dict = df.iloc[row_index].to_dict()
            
            # Replace text placeholders
            doc = self.replace_placeholders(doc, data_dict)
            
            # Embed images if image_handler is available
            if self.image_handler:
                doc = self.embed_images(doc, data_dict)
            
            # Save to BytesIO
            output = BytesIO()
            doc.save(output)
            output.seek(0)
            
            return output
        
        except Exception as e:
            raise Exception(f"Error generating document from template: {str(e)}")
    
    def convert_docx_to_pdf(self, docx_bytes, output_filename=None):
        """
        Convert DOCX to PDF.
        
        Args:
            docx_bytes (BytesIO): DOCX file as bytes
            output_filename (str): Optional output filename
            
        Returns:
            bytes: PDF file as bytes
        """
        try:
            # Initialize COM for Windows (required for docx2pdf)
            import pythoncom
            pythoncom.CoInitialize()
            
            try:
                # Create temporary files
                temp_docx = os.path.join(self.temp_dir, 'temp_report.docx')
                temp_pdf = os.path.join(self.temp_dir, 'temp_report.pdf')
                
                # Write DOCX to temp file
                with open(temp_docx, 'wb') as f:
                    f.write(docx_bytes.read())
                
                # Convert to PDF using docx2pdf (requires MS Word on Windows)
                convert(temp_docx, temp_pdf)
                
                # Read PDF bytes
                with open(temp_pdf, 'rb') as f:
                    pdf_bytes = f.read()
                
                # Cleanup temp files
                if os.path.exists(temp_docx):
                    os.remove(temp_docx)
                if os.path.exists(temp_pdf):
                    os.remove(temp_pdf)
                
                return pdf_bytes
            
            finally:
                # Always uninitialize COM
                pythoncom.CoUninitialize()
        
        except Exception as e:
            raise Exception(f"Error converting DOCX to PDF: {str(e)}. Make sure Microsoft Word is installed on Windows.")
    
    def generate_client_report(self, template_file, df, row_index=0):
        """
        Generate a client PDF report from DOCX template and data.
        
        Args:
            template_file: File-like object or path to template DOCX
            df (pd.DataFrame): DataFrame containing data
            row_index (int): Index of row to use for data
            
        Returns:
            bytes: PDF file as bytes
        """
        try:
            # Generate filled DOCX
            filled_docx = self.generate_from_template(template_file, df, row_index)
            
            # Convert to PDF
            pdf_bytes = self.convert_docx_to_pdf(filled_docx)
            
            return pdf_bytes
            
        except Exception as e:
            raise Exception(f"Error generating client report: {str(e)}")
    
    def generate_multiple_client_reports(self, template_file, df):
        """
        Generate multiple client PDF reports (one per row) and package them in a ZIP file.
        
        Args:
            template_file: File-like object or path to template DOCX
            df (pd.DataFrame): DataFrame containing data
            
        Returns:
            tuple: (zip_bytes, report_count, file_list)
                - zip_bytes: ZIP file as bytes containing all PDFs
                - report_count: Number of reports generated
                - file_list: List of generated filenames
        """
        import zipfile
        from io import BytesIO
        from datetime import datetime
        import re
        
        try:
            # Validate inputs
            if df is None:
                raise ValueError("DataFrame is None. Please ensure data is loaded and filtered correctly.")
            
            if not hasattr(df, '__len__'):
                raise ValueError(f"Invalid DataFrame type: {type(df)}. Expected pandas DataFrame.")
            
            if len(df) == 0:
                raise ValueError("DataFrame is empty. No records to generate reports for.")
            
            print(f"DEBUG: Starting generation for {len(df)} records")
            
            zip_buffer = BytesIO()
            file_list = []
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for index in range(len(df)):
                    print(f"DEBUG: Processing record {index + 1}/{len(df)}")
                    
                    # Get site name for filename
                    site_name = df.iloc[index].get('Site Name', f'Site_{index+1}')
                    print(f"DEBUG: Site name: {site_name}")
                    
                    # Clean site name for filename (remove invalid characters)
                    clean_site_name = re.sub(r'[<>:"/\\|?*]', '_', str(site_name))
                    clean_site_name = clean_site_name.strip()
                    
                    # Get date for filename
                    date_str = df.iloc[index].get('Date', datetime.now().strftime('%Y-%m-%d'))
                    if pd.notna(date_str):
                        date_str = str(date_str).replace('/', '-')
                    else:
                        date_str = datetime.now().strftime('%Y-%m-%d')
                    
                    # Generate filename
                    filename = f"{clean_site_name}_{date_str}_Report.pdf"
                    
                    # Avoid duplicate filenames
                    original_filename = filename
                    counter = 1
                    while filename in file_list:
                        filename = f"{clean_site_name}_{date_str}_Report_{counter}.pdf"
                        counter += 1
                    
                    file_list.append(filename)
                    print(f"DEBUG: Filename: {filename}")
                    
                    # Reset template file pointer
                    if hasattr(template_file, 'seek'):
                        template_file.seek(0)
                    
                    print(f"DEBUG: About to generate PDF for record {index}")
                    # Generate PDF for this record
                    try:
                        pdf_bytes = self.generate_client_report(template_file, df, row_index=index)
                        print(f"DEBUG: PDF generated successfully for record {index}")
                    except Exception as pdf_error:
                        print(f"DEBUG: Error in generate_client_report for record {index}: {str(pdf_error)}")
                        raise
                    
                    # Add to ZIP
                    zip_file.writestr(filename, pdf_bytes)
                    print(f"DEBUG: Added {filename} to ZIP")
            
            zip_buffer.seek(0)
            print(f"DEBUG: ZIP generation complete")
            return zip_buffer.getvalue(), len(df), file_list
        
        except Exception as e:
            # Provide detailed error information
            print(f"DEBUG: Exception caught: {type(e).__name__}: {str(e)}")
            import traceback
            traceback.print_exc()
            
            error_msg = f"Error generating multiple client reports: {str(e)}"
            if df is None:
                error_msg += " | DataFrame is None"
            elif not hasattr(df, '__len__'):
                error_msg += f" | DataFrame type issue: {type(df)}"
            raise Exception(error_msg)
